import os
import shutil
import json
from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import List, Optional

from .database import engine, Base, get_db
from . import models, schemas, document_proc, ai_engine, rag_engine

# Initialize database tables
Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="AI Study Companion API",
    description="Backend API for Document Uploads, Summarization, Flashcards, Quizzes, and RAG Chat.",
    version="1.0.0"
)

# CORS configurations
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if os.environ.get("VERCEL") == "1":
    UPLOAD_DIR = "/tmp/uploads"
else:
    UPLOAD_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "instance", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.get("/")
def read_root():
    return {"message": "Welcome to AI Study Companion API. Go to /docs for Swagger API documentation."}

@app.post("/api/upload", response_model=schemas.BookResponse)
def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Upload a book/document (PDF), extract its metadata and save it to the DB.
    Also, triggers RAG background indexing.
    """
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported currently.")
        
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    
    # Save file to uploads folder
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
        
    # Process metadata using fitz
    try:
        meta = document_proc.extract_pdf_metadata(file_path)
    except Exception as e:
        os.remove(file_path)
        raise HTTPException(status_code=400, detail=f"Failed to read PDF metadata: {str(e)}")
        
    # Check if book already exists in DB
    existing_book = db.query(models.Book).filter(models.Book.title == meta["title"]).first()
    if existing_book:
        # Re-use existing path if upload is repeated
        return existing_book

    # Create DB entry
    db_book = models.Book(
        title=meta["title"],
        author=meta["author"],
        total_pages=meta["total_pages"],
        language=meta["language"],
        file_path=file_path
    )
    db.add(db_book)
    db.commit()
    db.refresh(db_book)
    
    # Run RAG indexing in the background
    background_tasks.add_task(rag_engine.index_book, file_path, db_book.id, db_book.total_pages)
    
    return db_book

@app.get("/api/books", response_model=List[schemas.BookResponse])
def get_books(db: Session = Depends(get_db)):
    """
    List all uploaded books.
    """
    return db.query(models.Book).all()

@app.get("/api/books/{book_id}", response_model=schemas.BookResponse)
def get_book(book_id: int, db: Session = Depends(get_db)):
    """
    Get a single book details.
    """
    book = db.query(models.Book).filter(models.Book.id == book_id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
    return book

@app.post("/api/books/{book_id}/process", response_model=schemas.SummaryResponse)
async def process_book_range(
    book_id: int,
    req: schemas.SummaryRequest,
    db: Session = Depends(get_db)
):
    """
    Generate summaries, main ideas, concepts, and actionable insights for a specific page range.
    """
    book = db.query(models.Book).filter(models.Book.id == book_id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
        
    # Validate page range
    if req.start_page < 1 or req.end_page > book.total_pages or req.start_page > req.end_page:
        raise HTTPException(status_code=400, detail=f"Invalid page range. Book has {book.total_pages} pages.")
    if req.end_page - req.start_page + 1 > 50:
        raise HTTPException(status_code=400, detail="Maximum range for a single request is 50 pages. Please select a smaller page range.")
        
    # Check if a summary for this range already exists
    existing_summary = db.query(models.Summary).filter(
        models.Summary.book_id == book_id,
        models.Summary.start_page == req.start_page,
        models.Summary.end_page == req.end_page
    ).first()
    if existing_summary:
        print("Database Cache Hit for Summary!")
        return existing_summary

    import time
    start = time.time()

    # Extract text from specific range
    extract_start = time.time()
    try:
        extracted_text = document_proc.extract_text_range(book.file_path, req.start_page, req.end_page)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract page range: {str(e)}")
        
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="The selected page range contains no extractable text.")
    print("PDF Extraction Time:", time.time() - extract_start)
        
    # Format fields that could be parsed as lists from JSON responses
    def format_field(val) -> str:
        if isinstance(val, list):
            return "\n".join(f"- {item}" if not str(item).strip().startswith("-") and not str(item).strip().startswith("*") else str(item) for item in val)
        return str(val) if val is not None else ""

    # Generate summary with Llama 3.1
    ai_start = time.time()
    summary_data = await ai_engine.generate_summary(extracted_text)
    print("AI Summary Time:", time.time() - ai_start)
    print("Total Processing Time:", time.time() - start)
    
    # Save to database
    db_summary = models.Summary(
        book_id=book_id,
        start_page=req.start_page,
        end_page=req.end_page,
        summary_text=format_field(summary_data.get("summary", "")),
        main_idea=format_field(summary_data.get("main_idea", "")),
        key_concepts=format_field(summary_data.get("key_concepts", "")),
        definitions=format_field(summary_data.get("definitions", "")),
        actionable_insights=format_field(summary_data.get("actionable_insights", ""))
    )
    
    db.add(db_summary)
    db.commit()
    db.refresh(db_summary)
    
    return db_summary

@app.get("/api/books/{book_id}/summaries", response_model=List[schemas.SummaryResponse])
def get_book_summaries(book_id: int, db: Session = Depends(get_db)):
    """
    Get all generated summaries for a book.
    """
    return db.query(models.Summary).filter(models.Summary.book_id == book_id).all()

@app.post("/api/books/{book_id}/flashcards", response_model=List[schemas.FlashcardResponse])
async def generate_book_flashcards(
    book_id: int,
    req: schemas.SummaryRequest,
    db: Session = Depends(get_db)
):
    """
    Generate study flashcards for a specific page range.
    """
    book = db.query(models.Book).filter(models.Book.id == book_id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
        
    # Validate page range
    if req.start_page < 1 or req.end_page > book.total_pages or req.start_page > req.end_page:
        raise HTTPException(status_code=400, detail=f"Invalid page range. Book has {book.total_pages} pages.")
    if req.end_page - req.start_page + 1 > 50:
        raise HTTPException(status_code=400, detail="Maximum range for a single request is 50 pages. Please select a smaller page range.")
        
    # Attempt to pull cached summary
    summary = db.query(models.Summary).filter(
        models.Summary.book_id == book_id,
        models.Summary.start_page == req.start_page,
        models.Summary.end_page == req.end_page
    ).first()
    
    if summary:
        print("Generating flashcards using cached summary...")
        text = f"""
        Summary: {summary.summary_text}
        Main Idea: {summary.main_idea}
        Key Concepts: {summary.key_concepts}
        Definitions: {summary.definitions}
        Actionable Insights: {summary.actionable_insights}
        """
    else:
        print("No cached summary found. Extracting text from PDF...")
        try:
            text = document_proc.extract_text_range(book.file_path, req.start_page, req.end_page)
        except Exception as e:
            raise HTTPException(status_code=400, detail=str(e))
        
    cards = await ai_engine.generate_flashcards(text)
    
    db_cards = []
    for card in cards:
        db_card = models.Flashcard(
            book_id=book_id,
            front=card.get("front", ""),
            back=card.get("back", "")
        )
        db.add(db_card)
        db_cards.append(db_card)
        
    db.commit()
    for card in db_cards:
        db.refresh(card)
        
    return db_cards

@app.get("/api/books/{book_id}/flashcards", response_model=List[schemas.FlashcardResponse])
def get_book_flashcards_history(book_id: int, db: Session = Depends(get_db)):
    """
    Retrieve all previously generated flashcards for a book.
    """
    return db.query(models.Flashcard).filter(models.Flashcard.book_id == book_id).all()

@app.post("/api/books/{book_id}/quiz", response_model=schemas.QuizResponse)
async def generate_book_quiz(
    book_id: int,
    page_req: schemas.SummaryRequest,
    quiz_req: schemas.QuizRequest,
    db: Session = Depends(get_db)
):
    """
    Generate interactive quiz questions from a page range.
    """
    book = db.query(models.Book).filter(models.Book.id == book_id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
        
    # Validate page range
    if page_req.start_page < 1 or page_req.end_page > book.total_pages or page_req.start_page > page_req.end_page:
        raise HTTPException(status_code=400, detail=f"Invalid page range. Book has {book.total_pages} pages.")
    if page_req.end_page - page_req.start_page + 1 > 50:
        raise HTTPException(status_code=400, detail="Maximum range for a single request is 50 pages. Please select a smaller page range.")
        
    # Attempt to pull cached summary
    summary = db.query(models.Summary).filter(
        models.Summary.book_id == book_id,
        models.Summary.start_page == page_req.start_page,
        models.Summary.end_page == page_req.end_page
    ).first()
    
    if summary:
        print("Generating quiz using cached summary...")
        text = f"""
        Summary: {summary.summary_text}
        Main Idea: {summary.main_idea}
        Key Concepts: {summary.key_concepts}
        Definitions: {summary.definitions}
        Actionable Insights: {summary.actionable_insights}
        """
    else:
        print("No cached summary found. Extracting text from PDF...")
        try:
            text = document_proc.extract_text_range(book.file_path, page_req.start_page, page_req.end_page)
        except Exception as e:
            raise HTTPException(status_code=400, detail=str(e))
        
    questions = await ai_engine.generate_quiz(text, difficulty=quiz_req.difficulty, num_questions=quiz_req.num_questions)
    
    db_quiz = models.Quiz(
        book_id=book_id,
        difficulty=quiz_req.difficulty,
        quiz_data=json.dumps(questions)
    )
    db.add(db_quiz)
    db.commit()
    db.refresh(db_quiz)
    
    return schemas.QuizResponse(
        id=db_quiz.id,
        book_id=db_quiz.book_id,
        difficulty=db_quiz.difficulty,
        questions=questions,
        created_at=db_quiz.created_at
    )

@app.post("/api/books/{book_id}/chat")
def chat_endpoint(
    book_id: int,
    req: schemas.ChatRequest,
    db: Session = Depends(get_db)
):
    """
    Ask questions to the book using the local RAG pipeline.
    """
    # Verify book exists
    book = db.query(models.Book).filter(models.Book.id == book_id).first()
    if not book:
        raise HTTPException(status_code=404, detail="Book not found")
        
    # Save user message to chat history
    db_user_msg = models.ChatHistory(book_id=book_id, role="user", content=req.message)
    db.add(db_user_msg)
    
    # Run chat response
    response_text = rag_engine.chat_with_book(req.message, book_id, req.history)
    
    # Save assistant message to chat history
    db_assistant_msg = models.ChatHistory(book_id=book_id, role="assistant", content=response_text)
    db.add(db_assistant_msg)
    
    db.commit()
    
    return {"response": response_text}

@app.get("/api/books/{book_id}/chat-history")
def get_chat_history(book_id: int, db: Session = Depends(get_db)):
    """
    Retrieve all conversation history with a book.
    """
    messages = db.query(models.ChatHistory).filter(models.ChatHistory.book_id == book_id).order_by(models.ChatHistory.created_at.asc()).all()
    return [{"role": msg.role, "content": msg.content, "created_at": msg.created_at} for msg in messages]

@app.post("/api/translate")
async def translate_text(req: dict):
    """
    On-demand translation of content.
    """
    text = req.get("text", "")
    target_lang = req.get("target_language", "English")
    if not text:
        raise HTTPException(status_code=400, detail="Missing text for translation.")
    translated = await ai_engine.translate_content(text, target_lang)
    return {"translated": translated}

class DocxExportRequest(schemas.BaseModel):
    title: str
    author: str
    start_page: int
    end_page: int
    language: str
    summary_text: str
    main_idea: str
    key_concepts: str
    definitions: str
    actionable_insights: str

@app.post("/api/export/docx")
def export_docx(req: DocxExportRequest):
    """
    Export study notes as a Word Document (.docx) based on the user's active view.
    """
    import io
    from fastapi.responses import StreamingResponse
    from docx import Document
    
    doc = Document()
    doc.add_heading(req.title, level=0)
    doc.add_paragraph(f"Author: {req.author}")
    doc.add_paragraph(f"Page Range: {req.start_page} - {req.end_page}")
    doc.add_paragraph(f"Language: {req.language}")
    doc.add_paragraph("")
    
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(req.summary_text)
    
    doc.add_heading("2. Core Takeaway / Main Idea", level=1)
    doc.add_paragraph(req.main_idea)
    
    doc.add_heading("3. Key Concepts", level=1)
    doc.add_paragraph(req.key_concepts)
    
    doc.add_heading("4. Important Definitions", level=1)
    doc.add_paragraph(req.definitions)
    
    doc.add_heading("5. Actionable Insights", level=1)
    doc.add_paragraph(req.actionable_insights)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    clean_title = "".join(c for c in req.title if c.isalnum() or c in (' ', '_')).rstrip()
    filename = f"{clean_title.replace(' ', '_')}_Pages_{req.start_page}_to_{req.end_page}.docx"
    
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

